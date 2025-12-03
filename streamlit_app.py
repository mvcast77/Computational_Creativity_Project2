import streamlit as st
import json
import os
import requests
from openai import OpenAI
from datetime import datetime
from io import BytesIO
import re

# -------------------------
# 1. User uploads/inputs story idea
# -------------------------

st.set_page_config(page_title="Dynamic Outline", page_icon="📝", layout="wide")

# Helper functions for document generation
def generate_filename_from_story(story_text):
    """Generate a clean filename from story idea."""
    # Take first 50 characters, clean it up
    summary = story_text[:50].strip()
    # Remove special characters
    summary = re.sub(r'[^\w\s-]', '', summary)
    # Replace spaces with underscores
    summary = re.sub(r'\s+', '_', summary)
    # Add timestamp
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"{summary}_{timestamp}"

def generate_story_title(story_text):
    """Generate a short title from story idea."""
    # Take first sentence or first 60 characters
    first_sentence = story_text.split('.')[0] if '.' in story_text else story_text[:60]
    return first_sentence.strip()

def create_pdf_document(title, story_summary, outline_text):
    """Create a PDF document with proper formatting."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                                rightMargin=inch, leftMargin=inch,
                                topMargin=inch, bottomMargin=inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            fontName='Times-Roman',
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        # Body style - double spaced
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['BodyText'],
            fontSize=12,
            fontName='Times-Roman',
            leading=24,  # Double spacing (2 * 12pt)
            alignment=TA_LEFT
        )
        
        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.2*inch))
        
        # Add story summary
        story.append(Paragraph(f"<b>Story Summary:</b> {story_summary}", body_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Add outline
        for line in outline_text.split('\n'):
            if line.strip():
                story.append(Paragraph(line, body_style))
                story.append(Spacer(1, 0.1*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    except ImportError:
        st.error("PDF generation requires reportlab. Install with: pip install reportlab")
        return None

def create_docx_document(title, story_summary, outline_text):
    """Create a DOCX document with proper formatting."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Set paragraph spacing for double-spacing
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
        
        # Add title
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        
        doc.add_paragraph()  # Blank line
        
        # Add story summary
        summary_para = doc.add_paragraph()
        summary_para.add_run('Story Summary: ').bold = True
        summary_para.add_run(story_summary)
        
        doc.add_paragraph()  # Blank line
        
        # Add outline
        for line in outline_text.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except ImportError:
        st.error("DOCX generation requires python-docx. Install with: pip install python-docx")
        return None

def create_txt_document(title, story_summary, outline_text):
    """Create a plain text document."""
    content = f"{title}\n{'='*len(title)}\n\n"
    content += f"Story Summary: {story_summary}\n\n"
    content += outline_text
    return content

# Dark mode toggle in sidebar
if 'dark_mode' not in st.session_state:
    st.session_state.dark_mode = False

with st.sidebar:
    st.title("⚙️ Settings")
    
    # Theme section
    st.subheader("🎨 Theme")
    
    dark_mode = st.toggle("🌙 Dark Mode", value=st.session_state.dark_mode)
    if dark_mode != st.session_state.dark_mode:
        st.session_state.dark_mode = dark_mode
        st.rerun()
    
    # Download section
    if st.session_state.get('outline_generated', False):
        st.divider()
        st.subheader("📥 Download Outline")
        
        download_format = st.selectbox(
            "Select format:",
            ["PDF", "DOCX", "TXT"],
            key="download_format_select"
        )
        
        # Placeholder for story_idea_text and outline data
        # In a real app, these would be populated from session_state after generation
        # For now, we'll use dummy data or rely on actual session state if available
        
        # Get story idea for title and filename
        story_text = st.session_state.get('story_idea_text', 'Story Outline')
        title = generate_story_title(story_text)
        filename_base = generate_filename_from_story(story_text)
        
        # Construct outline text
        act1_text = "\n".join([f"- {beat}" for beat in st.session_state.get('act1_beats', [])])
        act2_text = "\n".join([f"- {beat}" for beat in st.session_state.get('act2_beats', [])])
        act3_text = "\n".join([f"- {beat}" for beat in st.session_state.get('act3_beats', [])])
        outline_text = f"Act I - Setup\n{act1_text}\n\nAct II - Rising Action\n{act2_text}\n\nAct III - Climax & Resolution\n{act3_text}"
        
        story_summary = story_text[:200] + "..." if len(story_text) > 200 else story_text
        
        if download_format == "PDF":
            pdf_buffer = create_pdf_document(title, story_summary, outline_text)
            if pdf_buffer:
                st.download_button(
                    label="📄 Download PDF",
                    data=pdf_buffer,
                    file_name=f"{filename_base}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
        elif download_format == "DOCX":
            docx_buffer = create_docx_document(title, story_summary, outline_text)
            if docx_buffer:
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_buffer,
                    file_name=f"{filename_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
        else:  # TXT
            txt_content = create_txt_document(title, story_summary, outline_text)
            st.download_button(
                label="📄 Download TXT",
                data=txt_content,
            )

# Apply dark mode CSS
if st.session_state.dark_mode:
    st.markdown("""
        <style>
        /* Dark mode styles */
        .stApp {
            background-color: #0e1117;
            color: #fafafa;
        }
        
        /* Sidebar styling */
        [data-testid="stSidebar"] {
            background-color: #262730;
        }
        [data-testid="stSidebar"] * {
            color: #fafafa !important;
        }
        
        /* Headers */
        h1, h2, h3, h4, h5, h6 {
            color: #fafafa !important;
        }
        
        /* Text inputs and areas */
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea,
        .stNumberInput > div > div > input {
            background-color: #262730;
            color: #fafafa;
            border-color: #4a4a4a;
        }
        
        /* Buttons */
        .stButton > button {
            background-color: #262730;
            color: #fafafa;
            border-color: #4a4a4a;
        }
        .stButton > button:hover {
            background-color: #3a3a4a;
        }
        
        /* Expanders */
        .streamlit-expanderHeader {
            background-color: #262730;
            color: #fafafa;
        }
        
        /* Selectbox */
        .stSelectbox > div > div {
            background-color: #262730;
            color: #fafafa;
        }
        
        /* Labels */
        label {
            color: #fafafa !important;
        }
        
        /* Captions */
        .stCaptionContainer {
            color: #b0b0b0 !important;
        }
        
        /* All markdown text */
        div[data-testid="stMarkdownContainer"] p {
            color: #fafafa !important;
        }
        </style>
    """, unsafe_allow_html=True)

# Custom CSS for generate button - cerulean color
st.markdown("""
    <style>
    /* Generate button customization */
    button[kind="primary"] {
        background-color: #2596be !important;
        border-color: #2596be !important;
    }
    button[kind="primary"]:hover {
        background-color: #1e7a9a !important;
        border-color: #1e7a9a !important;
    }
    </style>
""", unsafe_allow_html=True)

st.title("📝 Story Outline Builder")

st.subheader("Describe Your Story Idea in the Text Box Below or Upload it")

st.caption("💡 Tip: Any changes you make to the story text will automatically be included in the next LLM prompt when you generate or regenerate.")

# Create two columns for story input and file upload
col1, col2 = st.columns([2, 1])

with col1:
    story_idea = st.text_area(
        "Enter the premise or partial story description:",
        placeholder="Example: A young botanist discovers a glowing plant in the forest...",
        height=200
    )
    # Store in session state for document generation
    if story_idea:
        st.session_state.story_idea_text = story_idea

with col2:
    uploaded_file = st.file_uploader(
        "Upload a PDF, TXT, or DOCX file:",
        type=["pdf", "txt", "docx"]
    )
    
    # Plot points selection - moved here to align with upload button
    plot_points_per_act = st.number_input(
        "Number of plot points per act:",
        min_value=2,
        max_value=6,
        value=3,
        step=1,
        help="Enter how many key beats/plot points you want for each act (2-6)"
    )

file_text = ""

if uploaded_file is not None:
    st.success(f"Uploaded: {uploaded_file.name}")

    # Read file based on type
    if uploaded_file.type == "text/plain":
        file_text = uploaded_file.read().decode("utf-8")

    elif uploaded_file.type == "application/pdf":
        import PyPDF2
        reader = PyPDF2.PdfReader(uploaded_file)
        file_text = "\n".join([page.extract_text() for page in reader.pages])

    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        import docx
        doc = docx.Document(uploaded_file)
        file_text = "\n".join([p.text for p in doc.paragraphs])

    st.text_area("Extracted Document Text:", value=file_text, height=200)

# -------------------------
# 2. Connect to LLM API to make the og outline
# -------------------------
 
#API KEY GOES HERE
api_key = "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"
#MODEL NAME GOES HERE
model = "gpt-4o" 

client = OpenAI(api_key=api_key) #might not need this line

# -------------------------
# 3. Function: Call LLM
# -------------------------

# def call_llm(prompt):
#     response = requests.get("http://127.0.0.1:8000/api/v1/methods/receive_result") #might need to change 0.0.0.0 to 127.0.0.1
#     print("Received: ")
#     print(response.text)
#     data = response.json()
#     return data["new_text"]
#     # print("THE FUNCTION THAT CALLS THE LLM GOES HERE!")

#FOR WHEN WE NO LONGER USE API RESPONSE; THIS IS FOR THE LLM CALL
def call_llm(prompt: str) -> str:
    completion = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a helpful story outline assistant. When creating stories, promote diversity and inclusive representation of characters across race, ethnicity, gender, sexual orientation, religion, creed, and ideology."},
            {"role": "user", "content": prompt},
        ],
        temperature=0.8,
    )
    return completion.choices[0].message.content

# -------------------------
# 4. Generate Visual Outline with Editable Sections
# -------------------------

# Initialize session state for storing outline sections and version history
if 'full_outline' not in st.session_state:
    st.session_state.full_outline = ""
if 'act1_beats' not in st.session_state:
    st.session_state.act1_beats = []
if 'act2_beats' not in st.session_state:
    st.session_state.act2_beats = []
if 'act3_beats' not in st.session_state:
    st.session_state.act3_beats = []
if 'outline_generated' not in st.session_state:
    st.session_state.outline_generated = False
# Version history: list of dicts with keys 'timestamp', 'outline', 'act1_beats', 'act2_beats', 'act3_beats'
if 'outline_versions' not in st.session_state:
    st.session_state.outline_versions = []
if 'selected_version_idx' not in st.session_state:
    st.session_state.selected_version_idx = None
if 'plot_points_per_act' not in st.session_state:
    st.session_state.plot_points_per_act = 3

def _construct_full_outline_from_beats():
    act1_text = "\n".join([f"- {b}" for b in st.session_state.get('act1_beats', [])])
    act2_text = "\n".join([f"- {b}" for b in st.session_state.get('act2_beats', [])])
    act3_text = "\n".join([f"- {b}" for b in st.session_state.get('act3_beats', [])])
    return f"Act I - Setup\n{act1_text}\n\nAct II - Rising Action\n{act2_text}\n\nAct III - Climax & Resolution\n{act3_text}"


def _update_version_labels():
    # enforce label format: store just timestamp; UI will add index
    for i, v in enumerate(st.session_state.outline_versions):
        ts = v.get('timestamp') or datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        st.session_state.outline_versions[i]['label'] = ts



if st.button("🎬 Generate Complete Story Outline", type="primary"):
    if not story_idea and not file_text:
        st.error("Please enter a story idea or upload a document.")
    else:
        # Store the plot points preference
        st.session_state.plot_points_per_act = plot_points_per_act
        
        combined_text = (story_idea or "") + "\n\n" + (file_text or "")
        
        # Create example beats based on the selected number
        example_beats = "\n".join([f"                - Key beat {i+1}" for i in range(plot_points_per_act)])

        prompt = f"""
            Create a clear, detailed, visual story outline based on the following material:

            {combined_text}

            Your outline must follow this format with EXACTLY {plot_points_per_act} plot points per act:

            - Act I
            - Setup
{example_beats}
            - Act II
            - Rising Action
{example_beats}
            - Act III
            - Climax & Resolution
{example_beats}

            Guidelines:
            - Use hierarchical bullet formatting.
            - Provide EXACTLY {plot_points_per_act} key beats for each act.
            - Do NOT include a summary or explanations.
            - Keep the outline tight and structured like a screenplay or novel planner.
            - Focus purely on plot beats and story flow.
            - Promote diversity in characters: include diverse representation across race, ethnicity, gender, sexual orientation, religion, creed, and ideology.
            """

        with st.spinner("Generating outline..."):
            print(prompt)
            result = call_llm(prompt)

            #FOR API CALL
            # st.subheader("📘 Generated Story Outline")
            # st.markdown(combined_text)

            # Save current version before overwriting (if any outline exists)
            if st.session_state.full_outline.strip():
                ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                st.session_state.outline_versions.append({
                    'timestamp': ts,
                    'outline': st.session_state.full_outline,
                    'act1_beats': st.session_state.act1_beats[:],
                    'act2_beats': st.session_state.act2_beats[:],
                    'act3_beats': st.session_state.act3_beats[:]
                })
                _update_version_labels()

            st.session_state.full_outline = result #stores outline into current session
            st.session_state.outline_generated = True #marks outline as generated

            # Parse the outline into individual beats
            def parse_beats(text):
                """Extract individual beats from outline text, filtering out headers."""
                lines = text.split('\n')
                beats = []
                structural_keywords = ['setup', 'rising action', 'climax', 'resolution', 'climax & resolution', 'climax and resolution']
                
                for line in lines:
                    line = line.strip()
                    # Skip empty lines
                    if not line:
                        continue
                    
                    # Skip lines that start with "- **" (bold headers)
                    if line.startswith('- **') and line.endswith('**'):
                        continue
                    
                    # Extract actual beat content (remove bullet markers)
                    if line.startswith('- '):
                        beat = line[2:].strip()
                    else:
                        beat = line
                    
                    # Skip structural headers (check after removing bullet)
                    if beat.lower() in structural_keywords:
                        continue
                    
                    # Remove "Key beat X:" prefix if present
                    if beat.lower().startswith('key beat'):
                        if ':' in beat:
                            beat = beat.split(':', 1)[1].strip()
                        else:
                            continue
                    
                    if beat:
                        beats.append(beat)
                return beats

            # Split into per-act raw text and parse into beats
            lines = result.split('\n')
            act1_text = []
            act2_text = []
            act3_text = []
            current_act = 0

            for line in lines:
                line_stripped = line.strip()
                line_lower = line_stripped.lower()
                if 'act i' in line_lower and 'act ii' not in line_lower and 'act iii' not in line_lower:
                    current_act = 1
                    continue
                elif 'act ii' in line_lower and 'act iii' not in line_lower:
                    current_act = 2
                    continue
                elif 'act iii' in line_lower:
                    current_act = 3
                    continue
                if current_act == 1:
                    act1_text.append(line)
                elif current_act == 2:
                    act2_text.append(line)
                elif current_act == 3:
                    act3_text.append(line)

            st.session_state.act1_beats = parse_beats('\n'.join(act1_text))
            st.session_state.act2_beats = parse_beats('\n'.join(act2_text))
            st.session_state.act3_beats = parse_beats('\n'.join(act3_text))
            

            # Cap beats to the specified number of plot points per act
            max_beats = st.session_state.plot_points_per_act
            st.session_state.act1_beats = st.session_state.act1_beats[:max_beats]
            st.session_state.act2_beats = st.session_state.act2_beats[:max_beats]
            st.session_state.act3_beats = st.session_state.act3_beats[:max_beats]

            # After generating, set selected_version_idx to None (current)
            st.session_state.selected_version_idx = None

        st.success("✅ Outline generated! You can now edit individual sections below.")
        st.rerun()

# -------------------------
# 5. Editable Outline Sections (Linear Layout) + Version History UI
# -------------------------

if st.session_state.outline_generated or st.session_state.outline_versions:
    st.subheader("📘 Edit Your Story Outline")
    
    # Create two columns: outline on left, controls on right
    col_outline, col_controls = st.columns([2, 1])
    
    with col_controls:
        # Regenerate Entire Outline Section (at top of right column)
        st.markdown("**🔄 Regenerate Entire Outline**")
        regenerate_prompt = st.text_area(
            "Regeneration Instructions:",
            placeholder="Example: Make the story darker and add more mystery elements...",
            height=150,
            key="regenerate_prompt_input",
            label_visibility="collapsed"
        )
        
        if st.button("🔄 Regenerate Complete Outline", type="secondary", use_container_width=True):
            if not regenerate_prompt.strip():
                st.error("Please provide instructions for regenerating the outline.")
            else:
                # Save a version before regenerating
                if st.session_state.full_outline.strip():
                    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    st.session_state.outline_versions.append({
                        'timestamp': ts,
                        'outline': st.session_state.full_outline,
                        'act1_beats': st.session_state.act1_beats[:],
                        'act2_beats': st.session_state.act2_beats[:],
                        'act3_beats': st.session_state.act3_beats[:]
                    })
                    _update_version_labels()
                # include current outline and user edits in the prompt
                current_outline_text = _construct_full_outline_from_beats()
                combined_text = (story_idea or "") + "\n\n" + (file_text or "") + "\n\nCurrent Outline:\n" + current_outline_text
                prompt = f"""
                    Create a clear, detailed, visual story outline based on the following material:

                    {combined_text}

                    Additional instructions for this regeneration:
                    {regenerate_prompt}

                    Your outline must follow this format:

                    - Act I
                    - Setup
                        - Key beat 1
                        - Key beat 2
                    - Act II
                    - Rising Action
                        - Key beat 1
                        - Key beat 2
                    - Act III
                    - Climax & Resolution
                        - Key beat 1
                        - Key beat 2

                    Guidelines:
                    - Use hierarchical bullet formatting.
                    - Do NOT include a summary or explanations.
                    - Keep the outline tight and structured like a screenplay or novel planner.
                    - Focus purely on plot beats and story flow.
                    - Promote diversity in characters: include diverse representation across race, ethnicity, gender, sexual orientation, religion, creed, and ideology.
                    """

                with st.spinner("Regenerating complete outline..."):
                    result = call_llm(prompt)
                    st.session_state.full_outline = result
                    
                    # Reuse the parse_beats function from the initial generation
                    def parse_beats(text):
                        """Extract individual beats from outline text, filtering out headers."""
                        lines = text.split('\n')
                        beats = []
                        
                        for line in lines:
                            line = line.strip()
                            # Skip empty lines
                            if not line:
                                continue
                            
                            # Skip structural headers (but be more specific)
                            line_lower = line.lower()
                            if line_lower in ['setup', 'rising action', 'climax', 'resolution', 'climax & resolution', 'climax and resolution']:
                                continue
                            if line.startswith('- **') and line.endswith('**'):
                                # This is a bold header like "- **Setup**"
                                continue
                            
                            # Extract actual beat content (remove bullet markers)
                            if line.startswith('- '):
                                beat = line[2:].strip()
                                # Remove "Key beat X:" prefix if present
                                if beat.lower().startswith('key beat'):
                                    # Find the colon and take everything after it
                                    if ':' in beat:
                                        beat = beat.split(':', 1)[1].strip()
                                    else:
                                        continue  # Skip if it's just "Key beat X" without content
                                if beat:  # Only add non-empty beats
                                    beats.append(beat)
                        return beats
                    
                    # Parse the outline into sections
                    lines = result.split('\n')
                    act1_text = []
                    act2_text = []
                    act3_text = []
                    current_act = 0

                    for line in lines:
                        line_stripped = line.strip()
                        line_lower = line_stripped.lower()
                        
                        # Check for Act transitions - be more flexible
                        if 'act i' in line_lower and 'act ii' not in line_lower and 'act iii' not in line_lower:
                            current_act = 1
                            continue  # Don't add the act header line itself
                        elif 'act ii' in line_lower and 'act iii' not in line_lower:
                            current_act = 2
                            continue  # Don't add the act header line itself
                        elif 'act iii' in line_lower:
                            current_act = 3
                            continue  # Don't add the act header line itself
                        
                        # Assign line to current act
                        if current_act == 1:
                            act1_text.append(line)
                        elif current_act == 2:
                            act2_text.append(line)
                        elif current_act == 3:
                            act3_text.append(line)

                    st.session_state.act1_beats = parse_beats('\n'.join(act1_text))
                    st.session_state.act2_beats = parse_beats('\n'.join(act2_text))
                    st.session_state.act3_beats = parse_beats('\n'.join(act3_text))
                    
                st.success("✅ Outline regenerated successfully!")
                st.rerun()
        
        # Version History (below regenerate section)
        st.divider()
        st.markdown("**📚 Version History**")
        version_options = [f"{i+1}: {v.get('label', v.get('timestamp',''))}" for i, v in enumerate(st.session_state.outline_versions)]
        if st.session_state.outline_generated:
            version_options.append("Current")

        selected = st.selectbox(
            "Select version:",
            version_options,
            index=len(version_options)-1 if version_options else 0,
            key="version_selectbox",
            label_visibility="collapsed"
        )

    # If user selects a previous version, load its content into the editors (read-only)
    if selected != "Current":
        # parse index from option like '1: label'
        try:
            idx = int(selected.split(":")[0].strip()) - 1
        except Exception:
            idx = 0
        v = st.session_state.outline_versions[idx]
        st.info(f"Viewing version {idx+1}: {v.get('label', v.get('timestamp'))}. To restore, click below.")

        if st.button("Restore This Version", key=f"restore_{idx}"):
            st.session_state.full_outline = v['outline']
            st.session_state.act1_beats = v.get('act1_beats', [])
            st.session_state.act2_beats = v.get('act2_beats', [])
            st.session_state.act3_beats = v.get('act3_beats', [])
            st.session_state.outline_generated = True
            st.session_state.selected_version_idx = None
            st.success("Restored selected version. You can now edit and save as a new version.")
            st.rerun()

        # Show the outline sections as read-only
        with st.expander("📖 Act I - Setup", expanded=True):
            for i, beat in enumerate(v.get('act1_beats', [])):
                st.text_area(f"Beat {i+1}:", value=beat, height=80, key=f"act1_beat{i}_view_{idx}", disabled=True, label_visibility="collapsed")
        with st.expander("🎬 Act II - Rising Action", expanded=True):
            for i, beat in enumerate(v.get('act2_beats', [])):
                st.text_area(f"Beat {i+1}:", value=beat, height=80, key=f"act2_beat{i}_view_{idx}", disabled=True, label_visibility="collapsed")
        with st.expander("🎯 Act III - Climax & Resolution", expanded=True):
            for i, beat in enumerate(v.get('act3_beats', [])):
                st.text_area(f"Beat {i+1}:", value=beat, height=80, key=f"act3_beat{i}_view_{idx}", disabled=True, label_visibility="collapsed")
        st.stop()

    # Custom CSS to increase font size in text areas
    st.markdown("""
        <style>
        .stTextArea textarea {
            font-size: 16px !important;
            line-height: 1.6 !important;
        }
        </style>
    """, unsafe_allow_html=True)

    with col_outline:
        # Act I - Setup
        with st.expander("📖 Act I - Setup", expanded=True):
            for i in range(len(st.session_state.act1_beats)):
                st.session_state.act1_beats[i] = st.text_area(
                    f"Beat {i+1}:",
                    value=st.session_state.act1_beats[i],
                    height=80,
                    key=f"act1_beat_{i}",
                    label_visibility="collapsed"
                )
            # Regenerate Act I button
            if st.button("🔄 Regenerate Act I", key="regen_act1"):
                if st.session_state.full_outline.strip():
                    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    st.session_state.outline_versions.append({
                        'timestamp': ts,
                        'outline': st.session_state.full_outline,
                        'act1_beats': st.session_state.act1_beats[:],
                        'act2_beats': st.session_state.act2_beats[:],
                        'act3_beats': st.session_state.act3_beats[:]
                    })
                    _update_version_labels()

                current_outline_text = _construct_full_outline_from_beats()
                combined_text = (story_idea or "") + "\n\n" + (file_text or "") + "\n\nCurrent Outline:\n" + current_outline_text
                prompt = f"""Based on this story and the current outline: {combined_text}
                Generate ONLY Act I - Setup section with {st.session_state.plot_points_per_act} key beats in hierarchical bullet format.
                Promote diversity in characters: include diverse representation across race, ethnicity, gender, sexual orientation, religion, creed, and ideology."""
                with st.spinner("Regenerating Act I..."):
                    new_text = call_llm(prompt)
                def _parse_beats(text):
                    beats = []
                    for line in text.split('\n'):
                        l = line.strip()
                        if not l:
                            continue
                        if l.startswith('- '):
                            b = l[2:].strip()
                        else:
                            b = l
                        if b.lower().startswith('key beat') and ':' in b:
                            b = b.split(':',1)[1].strip()
                        if b:
                            beats.append(b)
                    return beats
                st.session_state.act1_beats = _parse_beats(new_text)
                st.session_state.act1_beats = st.session_state.act1_beats[:st.session_state.plot_points_per_act]
                st.rerun()

        # Act II - Rising Action
        with st.expander("🎬 Act II - Rising Action", expanded=True):
            for i in range(len(st.session_state.act2_beats)):
                st.session_state.act2_beats[i] = st.text_area(
                    f"Beat {i+1}:",
                    value=st.session_state.act2_beats[i],
                    height=80,
                    key=f"act2_beat_{i}",
                    label_visibility="collapsed"
                )
            # Regenerate Act II button
            if st.button("🔄 Regenerate Act II", key="regen_act2"):
                if st.session_state.full_outline.strip():
                    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    st.session_state.outline_versions.append({
                        'timestamp': ts,
                        'outline': st.session_state.full_outline,
                        'act1_beats': st.session_state.act1_beats[:],
                        'act2_beats': st.session_state.act2_beats[:],
                        'act3_beats': st.session_state.act3_beats[:]
                    })
                    _update_version_labels()

                current_outline_text = _construct_full_outline_from_beats()
                combined_text = (story_idea or "") + "\n\n" + (file_text or "") + "\n\nCurrent Outline:\n" + current_outline_text
                prompt = f"""Based on this story and the current outline: {combined_text}
                Generate ONLY Act II - Rising Action section with {st.session_state.plot_points_per_act} key beats in hierarchical bullet format.
                Promote diversity in characters: include diverse representation across race, ethnicity, gender, sexual orientation, religion, creed, and ideology."""
                with st.spinner("Regenerating Act II..."):
                    new_text = call_llm(prompt)
                def _parse_beats(text):
                    beats = []
                    for line in text.split('\n'):
                        l = line.strip()
                        if not l:
                            continue
                        if l.startswith('- '):
                            b = l[2:].strip()
                        else:
                            b = l
                        if b.lower().startswith('key beat') and ':' in b:
                            b = b.split(':',1)[1].strip()
                        if b:
                            beats.append(b)
                    return beats
                st.session_state.act2_beats = _parse_beats(new_text)
                st.session_state.act2_beats = st.session_state.act2_beats[:st.session_state.plot_points_per_act]
                st.rerun()

        # Act III - Climax & Resolution
        with st.expander("🎯 Act III - Climax & Resolution", expanded=True):
            for i in range(len(st.session_state.act3_beats)):
                st.session_state.act3_beats[i] = st.text_area(
                    f"Beat {i+1}:",
                    value=st.session_state.act3_beats[i],
                    height=80,
                    key=f"act3_beat_{i}",
                    label_visibility="collapsed"
                )
            # Regenerate Act III button
            if st.button("🔄 Regenerate Act III", key="regen_act3"):
                if st.session_state.full_outline.strip():
                    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                    st.session_state.outline_versions.append({
                        'timestamp': ts,
                        'outline': st.session_state.full_outline,
                        'act1_beats': st.session_state.act1_beats[:],
                        'act2_beats': st.session_state.act2_beats[:],
                        'act3_beats': st.session_state.act3_beats[:]
                    })
                    _update_version_labels()

                current_outline_text = _construct_full_outline_from_beats()
                combined_text = (story_idea or "") + "\n\n" + (file_text or "") + "\n\nCurrent Outline:\n" + current_outline_text
                prompt = f"""Based on this story and the current outline: {combined_text}
                Generate ONLY Act III - Climax & Resolution section with {st.session_state.plot_points_per_act} key beats in hierarchical bullet format.
                Promote diversity in characters: include diverse representation across race, ethnicity, gender, sexual orientation, religion, creed, and ideology."""
                with st.spinner("Regenerating Act III..."):
                    new_text = call_llm(prompt)
                def _parse_beats(text):
                    beats = []
                    for line in text.split('\n'):
                        l = line.strip()
                        if not l:
                            continue
                        if l.startswith('- '):
                            b = l[2:].strip()
                        else:
                            b = l
                        if b.lower().startswith('key beat') and ':' in b:
                            b = b.split(':',1)[1].strip()
                        if b:
                            beats.append(b)
                    return beats
                st.session_state.act3_beats = _parse_beats(new_text)
                st.session_state.act3_beats = st.session_state.act3_beats[:st.session_state.plot_points_per_act]
                st.rerun()



# -------------------------
# 6. Auto-Save and Clear Options
# -------------------------
    st.divider()
    
    # Auto-save the outline whenever beats are edited
    # Reconstruct full outline from beats automatically
    act1_text = "\n".join([f"- {beat}" for beat in st.session_state.act1_beats])
    act2_text = "\n".join([f"- {beat}" for beat in st.session_state.act2_beats])
    act3_text = "\n".join([f"- {beat}" for beat in st.session_state.act3_beats])
    st.session_state.full_outline = f"""Act I - Setup\n{act1_text}\n\nAct II - Rising Action\n{act2_text}\n\nAct III - Climax & Resolution\n{act3_text}"""
    
    # Clear outline button
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("🗑️ Clear Outline", use_container_width=True):
            st.session_state.outline_generated = False
            st.session_state.full_outline = ""
            st.session_state.act1_beats = []
            st.session_state.act2_beats = []
            st.session_state.act3_beats = []
            st.rerun()
